from __future__ import annotations

import math
import re
import uuid
from collections.abc import Sequence
from dataclasses import dataclass, field
from datetime import UTC, datetime
from io import BytesIO
from typing import Any, Awaitable, Callable, cast

import matplotlib
matplotlib.use("Agg")
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from matplotlib import pyplot as plt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as PdfImage
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.models import DataColumn, Dataset, Flag, PowerCurve, Project, TimeseriesData
from app.schemas.report import ReportColumnSelection, ReportFormat, ReportSectionId
from app.services.air_density import air_density_summary, calculate_air_density, estimate_pressure_from_elevation, monthly_averages, wind_power_density
from app.services.energy_estimate import ensure_seeded_default_power_curve, gross_energy_estimate, load_power_curve
from app.services.extreme_wind import extreme_wind_summary
from app.services.export_engine import ExportedArtifact
from app.services.qc_engine import get_clean_dataframe
from app.services.report_templates import DEFAULT_REPORT_SECTIONS, REPORT_SECTION_LABELS
from app.services.turbulence import calculate_ti, ti_by_speed_bin, ti_summary
from app.services.weibull import fit_weibull, weibull_pdf
from app.services.wind_shear import shear_profile


REPORT_BLUE = "#0f766e"
REPORT_GOLD = "#d97706"
REPORT_RED = "#dc2626"
REPORT_SLATE = "#334155"


@dataclass(slots=True)
class ReportImage:
    title: str
    image_bytes: bytes
    caption: str | None = None
    width_inches: float = 6.2


@dataclass(slots=True)
class ReportTable:
    title: str
    headers: list[str]
    rows: list[list[str]]


@dataclass(slots=True)
class ReportSection:
    section_id: ReportSectionId
    title: str
    paragraphs: list[str] = field(default_factory=list)
    tables: list[ReportTable] = field(default_factory=list)
    images: list[ReportImage] = field(default_factory=list)


@dataclass(slots=True)
class ReportContext:
    title: str
    subtitle: str
    generated_at: datetime
    project: Project
    dataset: Dataset
    row_count: int
    section_order: list[ReportSectionId]
    sections: list[ReportSection]


@dataclass(slots=True)
class ResolvedReportSelections:
    speed_column: DataColumn | None
    direction_column: DataColumn | None
    temperature_column: DataColumn | None
    pressure_column: DataColumn | None
    turbulence_column: DataColumn | None
    gust_column: DataColumn | None
    shear_speed_columns: list[DataColumn]
    power_curve: PowerCurve | None


def _slugify(value: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", value.strip().lower()).strip("-")
    return slug or "windwhisper-report"


def _report_file_name(project: Project, dataset: Dataset, report_format: ReportFormat) -> str:
    stem = f"{_slugify(project.name)}-{_slugify(dataset.name)}-report"
    return f"{stem}.{report_format}"


def _format_datetime(value: datetime | None) -> str:
    if value is None:
        return "Not available"
    return value.astimezone(UTC).strftime("%Y-%m-%d %H:%M UTC")


def _format_number(value: float | int | None, digits: int = 2, suffix: str = "") -> str:
    if value is None:
        return "Not available"
    return f"{float(value):,.{digits}f}{suffix}"


def _format_count(value: int | None) -> str:
    if value is None:
        return "0"
    return f"{value:,d}"


def _maybe_float(value: object) -> float | None:
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return float(value)
    if hasattr(value, "item"):
        try:
            return float(cast(Any, value).item())
        except (TypeError, ValueError):
            return None
    return None


def _find_first_column(dataset: Dataset, measurement_types: Sequence[str]) -> "DataColumn | None":
    for column in dataset.columns:
        if column.measurement_type in measurement_types:
            return column
    return None


def _resolve_column(dataset: Dataset, column_id: uuid.UUID, label: str) -> DataColumn:
    for column in dataset.columns:
        if column.id == column_id:
            return column
    raise ValueError(f"{label} does not belong to the selected dataset")


def _resolve_selected_column(
    dataset: Dataset,
    column_id: uuid.UUID | None,
    label: str,
    allowed_types: Sequence[str],
) -> DataColumn | None:
    if column_id is None:
        return _find_first_column(dataset, allowed_types)

    column = _resolve_column(dataset, column_id, label)
    if column.measurement_type not in set(allowed_types):
        raise ValueError(f"{label} must reference one of: {', '.join(allowed_types)}")
    return column


def _resolve_shear_columns(dataset: Dataset, shear_column_ids: list[uuid.UUID]) -> list[DataColumn]:
    if shear_column_ids:
        resolved_columns = [_resolve_column(dataset, column_id, "shear_column_ids") for column_id in shear_column_ids]
        if any(column.measurement_type != "speed" for column in resolved_columns):
            raise ValueError("shear_column_ids must reference wind speed columns")
        return sorted(resolved_columns, key=lambda item: (float(item.height_m or 0.0), item.name))

    return _speed_columns(dataset)


def _speed_columns(dataset: Dataset) -> list["DataColumn"]:
    return sorted(
        [column for column in dataset.columns if column.measurement_type == "speed" and column.height_m is not None],
        key=lambda item: (float(item.height_m or 0.0), item.name),
    )


def _coerce_numeric_series(frame: pd.DataFrame, column_name: str) -> pd.Series:
    if column_name not in frame.columns:
        return pd.Series(dtype=float)
    return pd.to_numeric(frame[column_name], errors="coerce").astype(float)


def _figure_bytes(width: float = 8.0, height: float = 4.6) -> tuple[Any, Any]:
    figure, axes = plt.subplots(figsize=(width, height), dpi=160)
    return figure, axes


def _save_figure(figure: Any) -> bytes:
    buffer = BytesIO()
    figure.tight_layout()
    figure.savefig(buffer, format="png", bbox_inches="tight", facecolor="white")
    plt.close(figure)
    return buffer.getvalue()


def _safe_table_rows(rows: Sequence[Sequence[str]]) -> list[list[str]]:
    return [[str(cell) for cell in row] for row in rows]


async def _load_project_dataset(db: AsyncSession, project_id: uuid.UUID, dataset_id: uuid.UUID) -> Dataset:
    statement = (
        select(Dataset)
        .options(
            selectinload(Dataset.project),
            selectinload(Dataset.columns),
            selectinload(Dataset.flags).selectinload(Flag.ranges),
            selectinload(Dataset.analysis_results),
        )
        .where(Dataset.id == dataset_id, Dataset.project_id == project_id)
    )
    dataset = (await db.execute(statement)).scalar_one_or_none()
    if dataset is None:
        raise ValueError("Dataset was not found for the selected project")
    return dataset


async def _count_rows(db: AsyncSession, dataset_id: uuid.UUID) -> int:
    count_value = await db.scalar(select(func.count(TimeseriesData.id)).where(TimeseriesData.dataset_id == dataset_id))
    return int(count_value or 0)


async def _resolve_power_curve(db: AsyncSession, power_curve_id: uuid.UUID | None) -> PowerCurve | None:
    if power_curve_id is None:
        return await ensure_seeded_default_power_curve(db)

    power_curve = await db.get(PowerCurve, power_curve_id)
    if power_curve is None:
        raise ValueError("Selected power curve was not found")
    return power_curve


async def _resolve_report_selections(
    db: AsyncSession,
    dataset: Dataset,
    column_selection: ReportColumnSelection | None,
    power_curve_id: uuid.UUID | None,
) -> ResolvedReportSelections:
    selection = column_selection or ReportColumnSelection()

    return ResolvedReportSelections(
        speed_column=_resolve_selected_column(dataset, selection.speed_column_id, "speed_column_id", ["speed"]),
        direction_column=_resolve_selected_column(dataset, selection.direction_column_id, "direction_column_id", ["direction"]),
        temperature_column=_resolve_selected_column(dataset, selection.temperature_column_id, "temperature_column_id", ["temperature"]),
        pressure_column=_resolve_selected_column(dataset, selection.pressure_column_id, "pressure_column_id", ["pressure"]),
        turbulence_column=_resolve_selected_column(dataset, selection.turbulence_column_id, "turbulence_column_id", ["speed_sd", "turbulence_intensity"]),
        gust_column=_resolve_selected_column(dataset, selection.gust_column_id, "gust_column_id", ["gust"]),
        shear_speed_columns=_resolve_shear_columns(dataset, selection.shear_column_ids),
        power_curve=await _resolve_power_curve(db, power_curve_id),
    )


async def _build_title_page(project: Project, dataset: Dataset, row_count: int, title: str | None) -> ReportSection:
    report_title = title or f"{project.name} Wind Resource Report"
    subtitle = f"Dataset: {dataset.name}"
    location = []
    if project.latitude is not None and project.longitude is not None:
        location.append(f"Lat/Lon {project.latitude:.5f}, {project.longitude:.5f}")
    if project.elevation is not None:
        location.append(f"Elevation {project.elevation:.0f} m")

    return ReportSection(
        section_id="title_page",
        title=REPORT_SECTION_LABELS["title_page"],
        paragraphs=[
            report_title,
            subtitle,
            "WindWhisper generated this report from the selected project dataset using the current QC exclusions and analysis defaults.",
            " | ".join(location) if location else "Location metadata not available",
            f"Rows analyzed: {_format_count(row_count)}",
            f"Generated: {_format_datetime(datetime.now(UTC))}",
        ],
    )


async def _build_executive_summary(
    db: AsyncSession,
    dataset: Dataset,
    exclude_flag_ids: list[uuid.UUID],
    selections: ResolvedReportSelections,
) -> ReportSection:
    speed_column = selections.speed_column
    if speed_column is None:
        return ReportSection(
            section_id="executive_summary",
            title=REPORT_SECTION_LABELS["executive_summary"],
            paragraphs=["No wind-speed column is available, so the executive summary could not calculate core wind statistics."],
        )

    frame = await get_clean_dataframe(db, dataset.id, column_ids=[speed_column.id], exclude_flag_ids=exclude_flag_ids)
    series = _coerce_numeric_series(frame, speed_column.name).dropna()
    if series.empty:
        return ReportSection(
            section_id="executive_summary",
            title=REPORT_SECTION_LABELS["executive_summary"],
            paragraphs=["All wind-speed samples were removed by QC filters or are missing in the selected dataset."],
        )

    positive = series.loc[series > 0.0]
    weibull_text = "Weibull fit unavailable"
    if positive.shape[0] >= 2:
        fit = fit_weibull(positive.to_numpy(dtype=float), method="mle")
        weibull_text = f"Weibull k {_format_number(float(fit['k']), 2)} and A {_format_number(float(fit['A']), 2, ' m/s')} with R² {_format_number(float(fit['r_squared']), 3)}"

    paragraphs = [
        f"Mean wind speed at {speed_column.name}: {_format_number(float(series.mean()), 2, ' m/s')}",
        f"Median wind speed: {_format_number(float(series.median()), 2, ' m/s')}; standard deviation {_format_number(float(series.std(ddof=0)), 2, ' m/s')}",
        weibull_text,
    ]

    power_curve_record = selections.power_curve
    if power_curve_record is not None:
        power_curve = load_power_curve({"points": power_curve_record.points_json or []})
        if len(series.index) >= 2:
            estimate = gross_energy_estimate(series.to_numpy(dtype=float), power_curve, timestamps=pd.DatetimeIndex(series.index))
            summary = estimate["summary"]
            paragraphs.append(
                f"Using the selected power curve {power_curve_record.name}, estimated annual gross energy is {_format_number(float(summary['annual_energy_mwh']), 1, ' MWh')} with capacity factor {_format_number(float(summary['capacity_factor_pct']), 1, '%')}."
            )

    return ReportSection(
        section_id="executive_summary",
        title=REPORT_SECTION_LABELS["executive_summary"],
        paragraphs=paragraphs,
    )


async def _build_site_description(project: Project, dataset: Dataset) -> ReportSection:
    metadata = dataset.metadata_json or {}
    paragraphs = [
        f"Project: {project.name}",
        f"Dataset source type: {dataset.source_type or 'Not specified'}; source file: {dataset.file_name or 'Not specified'}.",
        f"Latitude {_format_number(project.latitude, 5)}, longitude {_format_number(project.longitude, 5)}, elevation {_format_number(project.elevation, 0, ' m')}.",
    ]
    if metadata:
        metadata_parts = [f"{key}: {value}" for key, value in list(metadata.items())[:6]]
        paragraphs.append("Dataset metadata highlights: " + "; ".join(metadata_parts))

    table = ReportTable(
        title="Measured channels",
        headers=["Column", "Type", "Height (m)", "Unit"],
        rows=_safe_table_rows(
            [
                [column.name, column.measurement_type or "other", _format_number(column.height_m, 0), column.unit or ""]
                for column in dataset.columns
            ]
        ),
    )
    return ReportSection(
        section_id="site_description",
        title=REPORT_SECTION_LABELS["site_description"],
        paragraphs=paragraphs,
        tables=[table],
    )


async def _build_data_summary(db: AsyncSession, dataset: Dataset, row_count: int, exclude_flag_ids: list[uuid.UUID]) -> ReportSection:
    selected_column_ids = [column.id for column in dataset.columns]
    frame = await get_clean_dataframe(db, dataset.id, column_ids=selected_column_ids, exclude_flag_ids=exclude_flag_ids)
    completeness_rows: list[list[str]] = []
    for column in dataset.columns[:10]:
        clean_count = int(_coerce_numeric_series(frame, column.name).count())
        recovery_pct = (clean_count / row_count * 100.0) if row_count else 0.0
        completeness_rows.append([column.name, _format_count(clean_count), _format_number(recovery_pct, 1, "%")])

    paragraphs = [
        f"Time range: {_format_datetime(dataset.start_time)} to {_format_datetime(dataset.end_time)}.",
        f"Nominal time step: {_format_number(dataset.time_step_seconds, 0, ' s')}; total stored rows {_format_count(row_count)}.",
        f"Columns available: {_format_count(len(dataset.columns))}. QC exclusions applied: {_format_count(len(exclude_flag_ids))} flag selections.",
    ]

    table = ReportTable(
        title="Column recovery snapshot",
        headers=["Column", "Valid rows", "Recovery"],
        rows=completeness_rows,
    )
    return ReportSection(
        section_id="data_summary",
        title=REPORT_SECTION_LABELS["data_summary"],
        paragraphs=paragraphs,
        tables=[table],
    )


async def _build_qc_summary(dataset: Dataset, exclude_flag_ids: list[uuid.UUID]) -> ReportSection:
    if not dataset.flags:
        return ReportSection(
            section_id="qc_summary",
            title=REPORT_SECTION_LABELS["qc_summary"],
            paragraphs=["No QC flags have been defined for this dataset."],
        )

    rows: list[list[str]] = []
    excluded_set = set(exclude_flag_ids)
    for flag in dataset.flags:
        rows.append([
            flag.name,
            _format_count(len(flag.ranges)),
            "Excluded" if flag.id in excluded_set else "Included",
            flag.description or "",
        ])

    return ReportSection(
        section_id="qc_summary",
        title=REPORT_SECTION_LABELS["qc_summary"],
        paragraphs=[
            f"The dataset contains {_format_count(len(dataset.flags))} QC flags. Selected exclusions remove samples from the flagged ranges before analysis sections are calculated.",
        ],
        tables=[ReportTable(title="Flag inventory", headers=["Flag", "Ranges", "Status", "Description"], rows=rows)],
    )


async def _build_wind_rose(
    db: AsyncSession,
    dataset: Dataset,
    exclude_flag_ids: list[uuid.UUID],
    selections: ResolvedReportSelections,
) -> ReportSection:
    speed_column = selections.speed_column
    direction_column = selections.direction_column
    if speed_column is None or direction_column is None:
        return ReportSection(
            section_id="wind_rose",
            title=REPORT_SECTION_LABELS["wind_rose"],
            paragraphs=["Wind rose requires both one wind-speed and one wind-direction column."],
        )

    frame = await get_clean_dataframe(
        db,
        dataset.id,
        column_ids=[speed_column.id, direction_column.id],
        exclude_flag_ids=exclude_flag_ids,
    )
    rose = pd.DataFrame(
        {
            "speed": pd.to_numeric(frame[speed_column.name], errors="coerce"),
            "direction": pd.to_numeric(frame[direction_column.name], errors="coerce").mod(360.0),
        }
    ).dropna()
    if rose.empty:
        return ReportSection(
            section_id="wind_rose",
            title=REPORT_SECTION_LABELS["wind_rose"],
            paragraphs=["Wind rose could not be generated because no paired speed and direction samples remain after QC filtering."],
        )

    sector_width = 30.0
    sector_index = np.floor(((rose["direction"] + (sector_width / 2.0)) % 360.0) / sector_width).astype(int)
    frequencies = np.array([(sector_index == sector).mean() * 100.0 for sector in range(12)], dtype=float)
    theta = np.deg2rad(np.arange(0.0, 360.0, sector_width))

    figure = plt.figure(figsize=(6.4, 5.4), dpi=160)
    axis = cast(Any, figure.add_subplot(111, projection="polar"))
    axis.bar(theta, frequencies, width=np.deg2rad(sector_width * 0.92), color=REPORT_BLUE, alpha=0.85, edgecolor="white")
    axis.set_theta_zero_location("N")
    axis.set_theta_direction(-1)
    axis.set_title(f"Wind rose frequency by 30° sector\n{speed_column.name} / {direction_column.name}", fontsize=11, pad=18)

    top_sector = int(np.argmax(frequencies))
    paragraphs = [
        f"Prevailing sector center: {top_sector * 30:.0f}° with frequency {_format_number(float(frequencies[top_sector]), 1, '%')}.",
        f"Inputs used: speed column {speed_column.name} and direction column {direction_column.name}.",
    ]
    return ReportSection(
        section_id="wind_rose",
        title=REPORT_SECTION_LABELS["wind_rose"],
        paragraphs=paragraphs,
        images=[ReportImage(title="Wind rose", image_bytes=_save_figure(figure), caption="Frequency of wind occurrence by direction sector.")],
    )


async def _build_frequency_distribution(
    db: AsyncSession,
    dataset: Dataset,
    exclude_flag_ids: list[uuid.UUID],
    selections: ResolvedReportSelections,
) -> ReportSection:
    speed_column = selections.speed_column
    if speed_column is None:
        return ReportSection(
            section_id="frequency_distribution",
            title=REPORT_SECTION_LABELS["frequency_distribution"],
            paragraphs=["Frequency distribution requires a wind-speed column."],
        )

    frame = await get_clean_dataframe(db, dataset.id, column_ids=[speed_column.id], exclude_flag_ids=exclude_flag_ids)
    series = _coerce_numeric_series(frame, speed_column.name).dropna()
    if series.empty:
        return ReportSection(
            section_id="frequency_distribution",
            title=REPORT_SECTION_LABELS["frequency_distribution"],
            paragraphs=["No valid speed samples remain for the histogram and Weibull fit."],
        )

    positive = series.loc[series > 0.0]
    fit = fit_weibull(positive.to_numpy(dtype=float), method="mle") if positive.shape[0] >= 2 else None
    edges = np.histogram_bin_edges(series.to_numpy(dtype=float), bins=18)
    counts, bins = np.histogram(series.to_numpy(dtype=float), bins=edges)

    figure, axis = _figure_bytes()
    widths = np.diff(bins)
    total = max(int(counts.sum()), 1)
    axis.bar(bins[:-1], counts / total * 100.0, width=widths, align="edge", color="#cbd5e1", edgecolor="#475569", label="Observed")
    if fit is not None:
        x_values = np.linspace(float(bins[0]), float(bins[-1]), 180)
        pdf_values = weibull_pdf(x_values, float(fit["k"]), float(fit["A"]))
        representative_width = float(np.mean(widths)) if widths.size else 1.0
        axis.plot(x_values, pdf_values * representative_width * 100.0, color=REPORT_GOLD, linewidth=2.0, label="Weibull fit")
    axis.set_xlabel(f"Wind speed ({speed_column.unit or 'm/s'})")
    axis.set_ylabel("Frequency (%)")
    axis.set_title(f"Frequency distribution for {speed_column.name}")
    axis.legend(loc="upper right")

    paragraphs = [
        f"Mean wind speed {_format_number(float(series.mean()), 2, ' m/s')} and median {_format_number(float(series.median()), 2, ' m/s')}.",
    ]
    if fit is not None:
        paragraphs.append(
            f"Weibull fit recovered k {_format_number(float(fit['k']), 2)} and A {_format_number(float(fit['A']), 2, ' m/s')} with RMSE {_format_number(float(fit['rmse']), 3)}."
        )

    return ReportSection(
        section_id="frequency_distribution",
        title=REPORT_SECTION_LABELS["frequency_distribution"],
        paragraphs=paragraphs,
        images=[ReportImage(title="Histogram and Weibull fit", image_bytes=_save_figure(figure), caption="Observed wind-speed distribution with Weibull overlay.")],
    )


async def _build_wind_shear(
    db: AsyncSession,
    dataset: Dataset,
    exclude_flag_ids: list[uuid.UUID],
    selections: ResolvedReportSelections,
) -> ReportSection:
    speed_columns = selections.shear_speed_columns
    if len(speed_columns) < 2:
        return ReportSection(
            section_id="wind_shear",
            title=REPORT_SECTION_LABELS["wind_shear"],
            paragraphs=["Wind shear requires at least two wind-speed columns with distinct measurement heights."],
        )

    frame = await get_clean_dataframe(db, dataset.id, column_ids=[column.id for column in speed_columns], exclude_flag_ids=exclude_flag_ids)
    speeds_by_height = {
        float(column.height_m or 0.0): pd.to_numeric(frame[column.name], errors="coerce").to_numpy(dtype=float)
        for column in speed_columns
    }
    profile = shear_profile(speeds_by_height, column_ids_by_height={float(column.height_m or 0.0): column.id for column in speed_columns}, method="power")
    profile_points = profile.get("profile_points", [])
    pair_stats = profile.get("pair_stats", [])

    if not profile_points:
        return ReportSection(
            section_id="wind_shear",
            title=REPORT_SECTION_LABELS["wind_shear"],
            paragraphs=["Wind shear profile could not be derived from the available data."],
        )

    heights = [float(point["height_m"]) for point in profile_points if point.get("mean_speed") is not None]
    means = [float(point["mean_speed"]) for point in profile_points if point.get("mean_speed") is not None]
    figure, axis = _figure_bytes(width=5.8, height=4.8)
    axis.plot(means, heights, marker="o", linewidth=2.0, color=REPORT_BLUE)
    axis.set_xlabel("Mean speed (m/s)")
    axis.set_ylabel("Height (m)")
    axis.set_title("Vertical wind-speed profile")
    axis.grid(alpha=0.25)

    table_rows = [
        [
            f"{float(item['lower_height_m']):.0f}-{float(item['upper_height_m']):.0f}",
            _format_number(float(item.get("mean_alpha")) if item.get("mean_alpha") is not None else None, 3),
            _format_number(float(item.get("median_alpha")) if item.get("median_alpha") is not None else None, 3),
            _format_number(float(item.get("std_alpha")) if item.get("std_alpha") is not None else None, 3),
        ]
        for item in pair_stats
    ]

    representative = profile.get("representative_pair")
    paragraphs = []
    if representative is not None and representative.get("mean_alpha") is not None:
        paragraphs.append(f"Representative power-law exponent α is {_format_number(float(representative['mean_alpha']), 3)}.")
    paragraphs.append(f"Profile spans {len(profile_points)} measured heights from {min(heights):.0f} m to {max(heights):.0f} m.")

    return ReportSection(
        section_id="wind_shear",
        title=REPORT_SECTION_LABELS["wind_shear"],
        paragraphs=paragraphs,
        tables=[ReportTable(title="Pairwise shear statistics", headers=["Heights (m)", "Mean α", "Median α", "Std α"], rows=table_rows)],
        images=[ReportImage(title="Shear profile", image_bytes=_save_figure(figure), caption="Mean speed by sensor height.")],
    )


async def _build_turbulence(
    db: AsyncSession,
    dataset: Dataset,
    exclude_flag_ids: list[uuid.UUID],
    selections: ResolvedReportSelections,
) -> ReportSection:
    speed_column = selections.speed_column
    sd_column = selections.turbulence_column
    if speed_column is None or sd_column is None:
        return ReportSection(
            section_id="turbulence",
            title=REPORT_SECTION_LABELS["turbulence"],
            paragraphs=["Turbulence section requires one wind-speed column and one speed standard-deviation or TI column."],
        )

    frame = await get_clean_dataframe(db, dataset.id, column_ids=[speed_column.id, sd_column.id], exclude_flag_ids=exclude_flag_ids)
    speed_values = pd.to_numeric(frame[speed_column.name], errors="coerce").to_numpy(dtype=float)
    sd_values = pd.to_numeric(frame[sd_column.name], errors="coerce").to_numpy(dtype=float)
    ti_values = sd_values if sd_column.measurement_type == "turbulence_intensity" else calculate_ti(speed_values, sd_values)
    summary = ti_summary(speed_values, ti_values)
    speed_bins = ti_by_speed_bin(speed_values, ti_values, bin_width=1.0)
    if not speed_bins:
        return ReportSection(
            section_id="turbulence",
            title=REPORT_SECTION_LABELS["turbulence"],
            paragraphs=["Turbulence intensity could not be calculated from the selected columns."],
        )

    figure, axis = _figure_bytes()
    labels = [f"{float(item['lower']):g}-{float(item['upper']):g}" for item in speed_bins]
    values = [float(item["representative_ti"]) for item in speed_bins]
    axis.bar(labels, values, color=REPORT_BLUE, alpha=0.8)
    axis.set_ylabel("Representative TI")
    axis.set_xlabel("Wind speed bin (m/s)")
    axis.set_title("Representative turbulence intensity by speed bin")
    axis.tick_params(axis="x", rotation=45)

    paragraphs = [
        f"Mean turbulence intensity {_format_number(_maybe_float(summary.get('mean_ti')), 3)}.",
        f"Characteristic TI at 15 m/s {_format_number(_maybe_float(summary.get('characteristic_ti_15')), 3)} with suggested IEC class {summary.get('iec_class') or 'Not classified'}."
    ]
    return ReportSection(
        section_id="turbulence",
        title=REPORT_SECTION_LABELS["turbulence"],
        paragraphs=paragraphs,
        images=[ReportImage(title="TI by speed bin", image_bytes=_save_figure(figure), caption="Representative turbulence intensity across 1 m/s speed bins.")],
    )


async def _build_air_density(
    db: AsyncSession,
    dataset: Dataset,
    exclude_flag_ids: list[uuid.UUID],
    selections: ResolvedReportSelections,
) -> ReportSection:
    temperature_column = selections.temperature_column
    speed_column = selections.speed_column
    pressure_column = selections.pressure_column
    if temperature_column is None or speed_column is None:
        return ReportSection(
            section_id="air_density",
            title=REPORT_SECTION_LABELS["air_density"],
            paragraphs=["Air-density section requires at least temperature and wind-speed columns."],
        )

    column_ids = [temperature_column.id, speed_column.id]
    if pressure_column is not None:
        column_ids.append(pressure_column.id)
    frame = await get_clean_dataframe(db, dataset.id, column_ids=column_ids, exclude_flag_ids=exclude_flag_ids)
    if frame.empty:
        return ReportSection(
            section_id="air_density",
            title=REPORT_SECTION_LABELS["air_density"],
            paragraphs=["No valid samples remain for air-density calculations."],
        )

    temperature_values = pd.to_numeric(frame[temperature_column.name], errors="coerce").to_numpy(dtype=float)
    speed_values = pd.to_numeric(frame[speed_column.name], errors="coerce").to_numpy(dtype=float)
    if pressure_column is not None:
        pressure_values = pd.to_numeric(frame[pressure_column.name], errors="coerce").to_numpy(dtype=float)
        pressure_source = "measured"
        estimated_pressure_hpa = None
    elif dataset.project.elevation is not None:
        estimated_pressure_hpa = estimate_pressure_from_elevation(float(dataset.project.elevation))
        pressure_values = np.full(temperature_values.shape, estimated_pressure_hpa, dtype=float)
        pressure_source = "estimated"
    else:
        return ReportSection(
            section_id="air_density",
            title=REPORT_SECTION_LABELS["air_density"],
            paragraphs=["No pressure column or project elevation is available for air-density calculations."],
        )

    density_values = calculate_air_density(temperature_values, pressure_values)
    wpd_values = wind_power_density(speed_values, density_values)
    summary = air_density_summary(density_values, wpd_values)
    monthly = monthly_averages(pd.DatetimeIndex(frame.index), density_values, wpd_values)

    figure, axis = _figure_bytes()
    labels = [row["label"] for row in monthly]
    densities = [float(row["mean_density"]) for row in monthly]
    axis.plot(labels, densities, marker="o", color=REPORT_BLUE)
    axis.set_ylabel("Air density (kg/m³)")
    axis.set_title("Monthly mean air density")
    axis.tick_params(axis="x", rotation=45)

    paragraphs = [
        f"Pressure source used: {pressure_source}. Mean density {_format_number(_maybe_float(summary.get('mean_density')), 3, ' kg/m³')}.",
        f"Mean wind power density {_format_number(_maybe_float(summary.get('mean_wind_power_density')), 1, ' W/m²')}."
    ]
    if estimated_pressure_hpa is not None:
        paragraphs.append(f"Estimated pressure from elevation: {_format_number(estimated_pressure_hpa, 1, ' hPa')}.")

    return ReportSection(
        section_id="air_density",
        title=REPORT_SECTION_LABELS["air_density"],
        paragraphs=paragraphs,
        images=[ReportImage(title="Monthly air density", image_bytes=_save_figure(figure), caption="Monthly mean air density across the selected dataset.")],
    )


async def _build_extreme_wind(
    db: AsyncSession,
    dataset: Dataset,
    exclude_flag_ids: list[uuid.UUID],
    selections: ResolvedReportSelections,
) -> ReportSection:
    speed_column = selections.speed_column
    gust_column = selections.gust_column
    if speed_column is None:
        return ReportSection(
            section_id="extreme_wind",
            title=REPORT_SECTION_LABELS["extreme_wind"],
            paragraphs=["Extreme-wind analysis requires a wind-speed column."],
        )

    column_ids = [speed_column.id]
    if gust_column is not None:
        column_ids.append(gust_column.id)
    frame = await get_clean_dataframe(db, dataset.id, column_ids=column_ids, exclude_flag_ids=exclude_flag_ids)
    speed_series = _coerce_numeric_series(frame, speed_column.name).dropna()
    gust_series = _coerce_numeric_series(frame, gust_column.name).dropna() if gust_column is not None else None
    if speed_series.empty:
        return ReportSection(
            section_id="extreme_wind",
            title=REPORT_SECTION_LABELS["extreme_wind"],
            paragraphs=["No valid wind-speed samples remain for extreme-wind analysis."],
        )

    summary = extreme_wind_summary(speed_series, gust_series)
    curve = summary["return_period_curve"]
    figure, axis = _figure_bytes()
    axis.semilogx([row["return_period_years"] for row in curve], [row["wind_speed"] for row in curve], color=REPORT_RED, linewidth=2.0)
    axis.scatter([row["return_period_years"] for row in summary["observed_points"]], [row["wind_speed"] for row in summary["observed_points"]], color=REPORT_SLATE, s=20)
    axis.set_xlabel("Return period (years)")
    axis.set_ylabel("Wind speed (m/s)")
    axis.set_title("Extreme-wind return period fit")
    axis.grid(alpha=0.25)

    ve50 = None
    for item in summary["return_periods"]:
        if math.isclose(float(item["return_period_years"]), 50.0):
            ve50 = float(item["wind_speed"])
            break

    paragraphs = [
        f"Extreme-wind source: {summary['summary'].get('data_source', 'speed')}.",
        f"Estimated 50-year return wind speed {_format_number(ve50, 2, ' m/s')}.",
    ]

    annual_rows = [
        [str(row["year"]), _format_number(float(row["wind_speed"]), 2, ' m/s')]
        for row in summary["annual_maxima"]
    ]
    return ReportSection(
        section_id="extreme_wind",
        title=REPORT_SECTION_LABELS["extreme_wind"],
        paragraphs=paragraphs,
        tables=[ReportTable(title="Annual maxima", headers=["Year", "Wind speed"], rows=annual_rows)],
        images=[ReportImage(title="Extreme wind fit", image_bytes=_save_figure(figure), caption="Observed annual maxima against the fitted extreme-wind curve.")],
    )


async def _build_long_term_adjustment(dataset: Dataset) -> ReportSection:
    mcp_results = [item for item in dataset.analysis_results if item.analysis_type == "mcp"]
    if not mcp_results:
        return ReportSection(
            section_id="long_term_adjustment",
            title=REPORT_SECTION_LABELS["long_term_adjustment"],
            paragraphs=["No persisted MCP analysis result is available for this dataset. Run the MCP workspace and persist results before including this section in a final report."],
        )

    latest = sorted(mcp_results, key=lambda item: item.created_at or datetime.min.replace(tzinfo=UTC))[-1]
    results = latest.results or {}
    rows = [[str(key), str(value)] for key, value in list(results.items())[:10]]
    return ReportSection(
        section_id="long_term_adjustment",
        title=REPORT_SECTION_LABELS["long_term_adjustment"],
        paragraphs=["The latest stored MCP analysis result has been embedded below for traceability."],
        tables=[ReportTable(title="Stored MCP result snapshot", headers=["Key", "Value"], rows=rows)],
    )


async def _build_energy_estimate(
    db: AsyncSession,
    dataset: Dataset,
    exclude_flag_ids: list[uuid.UUID],
    selections: ResolvedReportSelections,
) -> ReportSection:
    speed_column = selections.speed_column
    if speed_column is None:
        return ReportSection(
            section_id="energy_estimate",
            title=REPORT_SECTION_LABELS["energy_estimate"],
            paragraphs=["Energy estimate requires a wind-speed column."],
        )

    curve_record = selections.power_curve
    if curve_record is None:
        return ReportSection(
            section_id="energy_estimate",
            title=REPORT_SECTION_LABELS["energy_estimate"],
            paragraphs=["No selected power curve is available, so the report cannot calculate an energy estimate."],
        )

    power_curve = load_power_curve({"points": curve_record.points_json or []})
    frame = await get_clean_dataframe(db, dataset.id, column_ids=[speed_column.id], exclude_flag_ids=exclude_flag_ids)
    speeds = _coerce_numeric_series(frame, speed_column.name).dropna()
    if speeds.empty or len(speeds.index) < 2:
        return ReportSection(
            section_id="energy_estimate",
            title=REPORT_SECTION_LABELS["energy_estimate"],
            paragraphs=["Not enough valid wind-speed samples remain for the gross energy estimate."],
        )

    estimate = gross_energy_estimate(speeds.to_numpy(dtype=float), power_curve, timestamps=pd.DatetimeIndex(speeds.index))
    summary = estimate["summary"]
    energy_frame = pd.DataFrame({"power_kw": estimate["power_kw"]}, index=pd.DatetimeIndex(speeds.index))
    energy_index = pd.DatetimeIndex(energy_frame.index)
    energy_frame["month"] = energy_index.month
    monthly = energy_frame.groupby("month")["power_kw"].mean().reset_index()

    figure, axis = _figure_bytes()
    axis.bar(monthly["month"].astype(str).tolist(), monthly["power_kw"].astype(float).tolist(), color=REPORT_GOLD, alpha=0.85)
    axis.set_xlabel("Month")
    axis.set_ylabel("Mean power (kW)")
    axis.set_title("Monthly mean power from the default turbine curve")

    paragraphs = [
        f"Power curve used: {curve_record.name}.",
        f"Estimated annual gross energy {_format_number(float(summary['annual_energy_mwh']), 1, ' MWh')}.",
        f"Capacity factor {_format_number(float(summary['capacity_factor_pct']), 1, '%')} with equivalent full-load hours {_format_number(float(summary['equivalent_full_load_hours']), 0, ' h')}."
    ]

    return ReportSection(
        section_id="energy_estimate",
        title=REPORT_SECTION_LABELS["energy_estimate"],
        paragraphs=paragraphs,
        images=[ReportImage(title="Monthly mean power", image_bytes=_save_figure(figure), caption="Monthly mean power derived from the default seeded turbine curve.")],
    )


async def _build_sections(
    db: AsyncSession,
    project: Project,
    dataset: Dataset,
    row_count: int,
    section_ids: list[ReportSectionId],
    exclude_flag_ids: list[uuid.UUID],
    title: str | None,
    selections: ResolvedReportSelections,
) -> list[ReportSection]:
    builders: dict[ReportSectionId, Callable[[], Awaitable[ReportSection]]] = {
        "title_page": lambda: _build_title_page(project, dataset, row_count, title),
        "executive_summary": lambda: _build_executive_summary(db, dataset, exclude_flag_ids, selections),
        "site_description": lambda: _build_site_description(project, dataset),
        "data_summary": lambda: _build_data_summary(db, dataset, row_count, exclude_flag_ids),
        "qc_summary": lambda: _build_qc_summary(dataset, exclude_flag_ids),
        "wind_rose": lambda: _build_wind_rose(db, dataset, exclude_flag_ids, selections),
        "frequency_distribution": lambda: _build_frequency_distribution(db, dataset, exclude_flag_ids, selections),
        "wind_shear": lambda: _build_wind_shear(db, dataset, exclude_flag_ids, selections),
        "turbulence": lambda: _build_turbulence(db, dataset, exclude_flag_ids, selections),
        "air_density": lambda: _build_air_density(db, dataset, exclude_flag_ids, selections),
        "extreme_wind": lambda: _build_extreme_wind(db, dataset, exclude_flag_ids, selections),
        "long_term_adjustment": lambda: _build_long_term_adjustment(dataset),
        "energy_estimate": lambda: _build_energy_estimate(db, dataset, exclude_flag_ids, selections),
    }

    sections: list[ReportSection] = []
    for section_id in section_ids:
        builder = builders[section_id]
        sections.append(await builder())
    return sections


def _render_docx(context: ReportContext) -> bytes:
    document = Document()
    normal_style = cast(Any, document.styles["Normal"])
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(10.5)

    for index, section in enumerate(context.sections):
        if index > 0:
            document.add_section(WD_SECTION_START.NEW_PAGE)

        if section.section_id == "title_page":
            heading = document.add_paragraph()
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = heading.add_run(section.paragraphs[0])
            run.bold = True
            run.font.size = Pt(22)
            for paragraph_text in section.paragraphs[1:]:
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run(paragraph_text)
            continue

        document.add_heading(section.title, level=1)
        for paragraph_text in section.paragraphs:
            document.add_paragraph(paragraph_text)

        for table_content in section.tables:
            document.add_paragraph(table_content.title).runs[0].bold = True
            table = document.add_table(rows=1, cols=len(table_content.headers))
            table.style = "Table Grid"
            for cell, header in zip(table.rows[0].cells, table_content.headers, strict=False):
                cell.text = header
            for row in table_content.rows:
                cells = table.add_row().cells
                for cell, value in zip(cells, row, strict=False):
                    cell.text = value

        for image in section.images:
            document.add_paragraph(image.title).runs[0].bold = True
            document.add_picture(BytesIO(image.image_bytes), width=Inches(image.width_inches))
            if image.caption:
                caption = document.add_paragraph(image.caption)
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _render_pdf(context: ReportContext) -> bytes:
    buffer = BytesIO()
    document = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=42, rightMargin=42, topMargin=42, bottomMargin=42)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("ReportTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=24, leading=28, textColor=colors.HexColor(REPORT_SLATE), alignment=1)
    heading_style = ParagraphStyle("ReportHeading", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=16, leading=20, textColor=colors.HexColor(REPORT_SLATE), spaceAfter=10)
    body_style = ParagraphStyle("ReportBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5, leading=15)
    caption_style = ParagraphStyle("ReportCaption", parent=styles["BodyText"], fontName="Helvetica-Oblique", fontSize=9, leading=12, alignment=1, textColor=colors.HexColor("#475569"))

    story: list[Any] = []
    for index, section in enumerate(context.sections):
        if index > 0:
            story.append(PageBreak())

        if section.section_id == "title_page":
            story.append(Spacer(1, 1.4 * inch))
            story.append(Paragraph(section.paragraphs[0], title_style))
            story.append(Spacer(1, 0.25 * inch))
            for paragraph_text in section.paragraphs[1:]:
                story.append(Paragraph(paragraph_text, ParagraphStyle("CenteredBody", parent=body_style, alignment=1)))
                story.append(Spacer(1, 0.12 * inch))
            continue

        story.append(Paragraph(section.title, heading_style))
        for paragraph_text in section.paragraphs:
            story.append(Paragraph(paragraph_text, body_style))
            story.append(Spacer(1, 0.09 * inch))

        for table_content in section.tables:
            story.append(Spacer(1, 0.08 * inch))
            story.append(Paragraph(table_content.title, ParagraphStyle("TableTitle", parent=body_style, fontName="Helvetica-Bold")))
            data = [table_content.headers, *table_content.rows]
            table = Table(data, repeatRows=1)
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(REPORT_BLUE)),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#cbd5e1")),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 0.14 * inch))

        for image in section.images:
            story.append(Paragraph(image.title, ParagraphStyle("ImageTitle", parent=body_style, fontName="Helvetica-Bold")))
            story.append(Spacer(1, 0.05 * inch))
            story.append(PdfImage(BytesIO(image.image_bytes), width=image.width_inches * inch, height=image.width_inches * inch * 0.62))
            if image.caption:
                story.append(Paragraph(image.caption, caption_style))
            story.append(Spacer(1, 0.12 * inch))

    document.build(story)
    return buffer.getvalue()


async def generate_report(
    db: AsyncSession,
    project_id: uuid.UUID,
    *,
    dataset_id: uuid.UUID,
    sections: list[ReportSectionId] | None,
    report_format: ReportFormat,
    exclude_flag_ids: list[uuid.UUID] | None = None,
    title: str | None = None,
    column_selection: ReportColumnSelection | None = None,
    power_curve_id: uuid.UUID | None = None,
) -> ExportedArtifact:
    dataset = await _load_project_dataset(db, project_id, dataset_id)
    project = dataset.project
    if project is None:
        raise ValueError("Selected dataset is missing its parent project")

    section_ids = list(sections or DEFAULT_REPORT_SECTIONS)
    exclude_ids = list(exclude_flag_ids or [])
    row_count = await _count_rows(db, dataset.id)
    selections = await _resolve_report_selections(db, dataset, column_selection, power_curve_id)
    built_sections = await _build_sections(db, project, dataset, row_count, section_ids, exclude_ids, title, selections)

    context = ReportContext(
        title=title or f"{project.name} Wind Resource Report",
        subtitle=dataset.name,
        generated_at=datetime.now(UTC),
        project=project,
        dataset=dataset,
        row_count=row_count,
        section_order=section_ids,
        sections=built_sections,
    )

    if report_format == "docx":
        content = _render_docx(context)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        content = _render_pdf(context)
        media_type = "application/pdf"

    return ExportedArtifact(
        content=content,
        file_name=_report_file_name(project, dataset, report_format),
        media_type=media_type,
    )
